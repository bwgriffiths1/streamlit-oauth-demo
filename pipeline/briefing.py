"""
briefing.py — Phase 2: render the meeting-level briefing as a .docx file.

Takes the raw markdown text produced by rollup.generate_meeting_briefing_text()
and writes a clean, lightly-formatted Word document:
  - Title: committee + meeting date(s)
  - Body: parsed markdown sections → Word headings / bullets / body text
  - Inline images: renders <!-- image_id:N --> comments as embedded figures

Two rendering modes:
  - "classic" — original plain Word styles (build_and_save_briefing)
  - "redesigned" — polished NEPOOL brand design (build_and_save_briefing_v2)

All functions return Path / data rather than printing. Safe to call from Streamlit.
"""
import base64
import logging
import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, Twips, RGBColor

logger = logging.getLogger(__name__)

# Pattern for inline image references embedded by the summarizer
_IMAGE_REF_RE = re.compile(r"<!--\s*image_id:(\d+)\s*-->")


def _fetch_image_record(image_id: int) -> dict | None:
    """Fetch an image record from DB by ID. Returns None on failure."""
    try:
        import pipeline.db_new as db
        rows = db.get_images_by_ids([image_id])
        return rows[0] if rows else None
    except Exception as exc:
        logger.warning("Failed to fetch image %d: %s", image_id, exc)
        return None


def _add_image_to_doc(doc: Document, image_id: int, width: float = 5.5) -> None:
    """
    Fetch an image by DB id and embed it inline in the Word document.
    Adds a caption paragraph below the image.
    """
    record = _fetch_image_record(image_id)
    if not record or not record.get("image_b64"):
        return

    img_bytes = base64.b64decode(record["image_b64"])
    caption = record.get("description") or ""

    # Write to temp file (python-docx needs a file path or file-like object)
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        tmp.write(img_bytes)
        tmp_path = tmp.name

    try:
        doc.add_picture(tmp_path, width=Inches(width))
    except Exception as exc:
        logger.warning("Failed to embed image %d in docx: %s", image_id, exc)
        return
    finally:
        try:
            Path(tmp_path).unlink()
        except OSError:
            pass

    if caption:
        cap_para = doc.add_paragraph()
        cap_para.paragraph_format.space_before = Pt(2)
        cap_para.paragraph_format.space_after = Pt(8)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_para.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ---------------------------------------------------------------------------
# Markdown → docx rendering
# ---------------------------------------------------------------------------

def _add_horizontal_rule(doc: Document) -> None:
    """Insert a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run_with_inline_markup(para, text: str) -> None:
    """
    Add runs to an existing paragraph, honouring **bold** inline markup.
    Everything else is added as plain text.
    """
    # Strip CommonMark backslash escapes (\$ → $, \- → -, \( → (, etc.)
    # fix_markdown() writes these into .md files so markdown renderers see
    # literal characters rather than LaTeX math delimiters.  python-docx
    # does not interpret CommonMark, so we un-escape before building Word runs.
    text = re.sub(r'\\(.)', r'\1', text)

    # Split on **...**
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        m = re.fullmatch(r"\*\*([^*]+)\*\*", part)
        if m:
            run = para.add_run(m.group(1))
            run.bold = True
        else:
            para.add_run(part)


def _render_markdown_to_docx(doc: Document, text: str) -> None:
    """
    Render Claude's markdown output into the Word document.

    Handles:
      ## Heading 2  → Word Heading 2
      ### Heading 3 → Word Heading 3
      - bullet      → Word List Bullet
      ---           → horizontal rule
      <!-- image_id:N --> → inline image from database
      blank line    → paragraph break (collapsed)
      **bold**      → inline bold
      everything else → Normal paragraph
    """
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            # Blank line — skip (Word adds its own spacing)
            i += 1
            continue

        # Inline image reference
        img_match = _IMAGE_REF_RE.search(stripped)
        if img_match:
            _add_image_to_doc(doc, int(img_match.group(1)))
            i += 1
            continue

        if stripped == "---":
            _add_horizontal_rule(doc)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("# "):
            # Top-level heading inside body — treat as H2 to keep hierarchy
            doc.add_heading(stripped[2:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_run_with_inline_markup(para, stripped[2:])
            i += 1
            continue

        # Plain paragraph (may span until the next blank line)
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        _add_run_with_inline_markup(para, stripped)
        i += 1


# ---------------------------------------------------------------------------
# Date formatting helper
# ---------------------------------------------------------------------------

def _format_date_range(iso_dates: list[str]) -> str:
    """
    Format a list of ISO date strings as a human-readable range.

    ["2026-03-10"]                   → "March 10, 2026"
    ["2026-03-10", "2026-03-12"]     → "March 10–12, 2026"
    ["2026-03-31", "2026-04-01"]     → "March 31–April 1, 2026"
    """
    from datetime import date

    if not iso_dates:
        return ""
    dates = sorted(date.fromisoformat(d) for d in iso_dates)
    if len(dates) == 1:
        return dates[0].strftime("%B %-d, %Y")
    first, last = dates[0], dates[-1]
    if first.month == last.month and first.year == last.year:
        return f"{first.strftime('%B %-d')}–{last.strftime('%-d, %Y')}"
    if first.year == last.year:
        return f"{first.strftime('%B %-d')}–{last.strftime('%B %-d, %Y')}"
    return f"{first.strftime('%B %-d, %Y')}–{last.strftime('%B %-d, %Y')}"


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_and_save_briefing(
    briefing_text: str,
    manifest: dict,
    meeting_folder: Path,
) -> Path:
    """
    Build the .docx briefing from Claude's markdown output and save it
    to meeting_folder/<folder_name> Briefing.docx.

    Returns the Path of the saved file.
    """
    doc = Document()

    # --- Page margins (1 inch all around) ---
    for section in doc.sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    # --- Title block ---
    committee = manifest.get("committee", "Committee")
    dates = manifest.get("meeting_dates", [])
    date_str = _format_date_range(dates)

    title_para = doc.add_heading(f"{committee}", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sub_para = doc.add_paragraph(f"Meeting Briefing  ·  {date_str}")
    sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_para.paragraph_format.space_after = Pt(12)
    if sub_para.runs:
        sub_para.runs[0].italic = True
        sub_para.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    _add_horizontal_rule(doc)

    # --- Body: Claude's markdown output ---
    _render_markdown_to_docx(doc, briefing_text)

    # --- Save ---
    folder_name = meeting_folder.name          # e.g. "MC March 10 2026"
    filename = f"{folder_name} Briefing.docx"
    output_path = meeting_folder / filename
    doc.save(str(output_path))

    logger.info("Briefing saved: %s", output_path)
    return output_path


# ---------------------------------------------------------------------------
# Redesigned briefing (v2) — NEPOOL brand design system
# ---------------------------------------------------------------------------

# Brand colors
_CHARCOAL   = RGBColor(0x3A, 0x3A, 0x3B)
_CYAN       = RGBColor(0x00, 0xAD, 0xEF)
_CYAN_BG    = "EAF7FD"
_GRAY_BG    = "F4F5F6"
_GRAY_MID   = RGBColor(0xC8, 0xCA, 0xCC)
_GRAY_TEXT  = RGBColor(0x88, 0x88, 0x88)
_CYAN_HEX     = "00ADEF"
_GRAY_MID_HEX = "C8CACC"
_CONTENT_W    = 9360  # 6.5" in twips


def _v2_run(para, text, *, size=Pt(10.5), bold=False, color=_CHARCOAL):
    r = para.add_run(text)
    r.font.name = "Calibri"; r.font.size = size
    r.bold = bold; r.font.color.rgb = color
    return r


def _v2_pborder(para, side, sz, color, space=0):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr"); pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), str(space)); el.set(qn("w:color"), color)
    pBdr.append(el)


def _v2_right_tab(para, pos=_CONTENT_W):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), str(pos))
    tabs.append(tab); pPr.append(tabs)


def _v2_cell_borders(cell, **sides):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side, attrs in sides.items():
        if attrs:
            el = OxmlElement(f"w:{side}")
            for k, v in attrs.items():
                el.set(qn(f"w:{k}"), str(v))
            tcB.append(el)
    tcPr.append(tcB)


def _v2_cell_shading(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)


def _v2_cell_margins(cell, *, top=0, bottom=0, left=0, right=0):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _v2_pshading(para, fill):
    """Apply background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _v2_pindent(para, *, left=0, right=0):
    """Set paragraph left/right indent (twips)."""
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    if left: ind.set(qn("w:left"), str(left))
    if right: ind.set(qn("w:right"), str(right))
    pPr.append(ind)


def _v2_spacing(para, *, before=None, after=None, line=None):
    pf = para.paragraph_format
    if before is not None: pf.space_before = before
    if after is not None: pf.space_after = after
    if line is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line


def _v2_bold_runs(para, text, *, size=Pt(10.5), color=_CHARCOAL):
    text = re.sub(r"\\(.)", r"\1", text)
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        m = re.fullmatch(r"\*\*([^*]+)\*\*", part)
        if m:
            _v2_run(para, m.group(1), size=size, bold=True, color=color)
        elif part:
            _v2_run(para, part, size=size, bold=False, color=color)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and len(s) > 2


def _is_separator_row(line: str) -> bool:
    """Return True for markdown separator lines like |---|---|"""
    s = line.strip()
    if not _is_table_row(s):
        return False
    return all(re.fullmatch(r"\s*:?-+:?\s*", cell) for cell in s[1:-1].split("|"))


def _parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip()[1:-1].split("|")]


def _add_word_table(doc: Document, rows: list[list[str]]) -> None:
    """Render a list of cell-lists as a Word table; first row is treated as header."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row_data in enumerate(rows):
        for ci in range(ncols):
            cell_text = row_data[ci] if ci < len(row_data) else ""
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _v2_bold_runs(p, cell_text, size=Pt(9.5), color=_CHARCOAL)
            if ri == 0:
                for run in p.runs:
                    run.bold = True


def _render_v2_body_lines(doc: Document, body_lines: list[str]) -> None:
    """Render body text lines, converting consecutive | rows into Word tables
    and <!-- image_id:N --> references into inline images."""
    i = 0
    while i < len(body_lines):
        line = body_lines[i]
        # Inline image reference
        img_match = _IMAGE_REF_RE.search(line)
        if img_match:
            _add_image_to_doc(doc, int(img_match.group(1)))
            i += 1
            continue
        # H4 sub-heading: italic blue number, bold italic black title
        if line.strip().startswith("#### "):
            h4_text = line.strip()[5:].strip()
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(12), after=Pt(4))
            if ":" in h4_text:
                num, title = h4_text.split(":", 1)
                r = p.add_run(num.strip())
                r.font.name = "Calibri"; r.font.size = Pt(10.5)
                r.italic = True; r.font.color.rgb = _CYAN
                r = p.add_run(":  " + title.strip())
                r.font.name = "Calibri"; r.font.size = Pt(10.5)
                r.bold = True; r.italic = True; r.font.color.rgb = _CHARCOAL
            else:
                r = p.add_run(h4_text)
                r.font.name = "Calibri"; r.font.size = Pt(10.5)
                r.bold = True; r.italic = True; r.font.color.rgb = _CHARCOAL
            i += 1
            continue
        if _is_table_row(line):
            table_lines: list[str] = []
            while i < len(body_lines) and _is_table_row(body_lines[i]):
                table_lines.append(body_lines[i])
                i += 1
            data_rows = [_parse_table_row(r) for r in table_lines if not _is_separator_row(r)]
            _add_word_table(doc, data_rows)
            doc.add_paragraph()  # spacer after table
        else:
            p = doc.add_paragraph()
            _v2_spacing(p, before=Pt(0), after=Pt(7), line=Pt(12.1))
            _v2_bold_runs(p, line)
            i += 1


def _v2_page_number(run):
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
    run._r.extend([fc1, it, fc2])


def _render_v2_exec_summary(doc: Document, exec_lines: list[str]) -> None:
    """
    Render the executive summary as a single shaded box (like NEXT STEPS).
    All lines in one paragraph with \\n separators so the cyan background
    is one continuous block with no gaps.
    """
    if not exec_lines:
        return

    p = doc.add_paragraph()
    _v2_pshading(p, _CYAN_BG)
    _v2_pborder(p, "top", 10, _CYAN_HEX, space=6)
    _v2_pborder(p, "bottom", 10, _CYAN_HEX, space=6)
    _v2_spacing(p, before=Pt(0), after=Pt(0), line=Pt(12.1))

    for i, line in enumerate(exec_lines):
        if i > 0:
            p.add_run("\n")
        # Handle bullet lines
        if line.startswith("- ") or line.startswith("* "):
            _v2_run(p, "–  ", size=Pt(10.5), color=_CHARCOAL)
            _v2_bold_runs(p, line[2:])
        else:
            _v2_bold_runs(p, line)


def _v2_parse_briefing_md(text: str) -> dict:
    """Parse Claude's briefing markdown into structured data for the v2 renderer."""
    lines = text.splitlines()
    data = {"title": "", "subtitle": "", "date": "", "exec_summary": [], "items": []}
    i = 0

    # H1 (optional — Claude's briefing prompt does not require one)
    # Use a for-loop so i is only advanced if an H1 is actually found;
    # if none exists the main content scan still starts from line 0.
    for j, line in enumerate(lines):
        s = line.strip()
        if s.startswith("## "):
            break  # hit a real section heading — no H1 present
        if s.startswith("# "):
            data["title"] = s[2:].strip()
            i = j + 1
            break

    # Subtitle / date (only scan a short window after the H1, if any)
    j = i
    while j < min(i + 5, len(lines)):
        s = lines[j].strip()
        if not s:
            j += 1
            continue
        if s.startswith("#"):
            break  # hit a section heading — no subtitle line
        m = re.match(r"^\*(.+)\*$", s)
        inner = m.group(1) if m else s
        if "·" in inner:
            parts = inner.split("·", 1)
            data["subtitle"] = parts[0].strip()
            data["date"] = parts[1].strip()
        else:
            data["subtitle"] = inner
        i = j + 1
        break
        j += 1

    exec_lines, cur = [], None
    mode = "scan"
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("## "):
            h = s[3:].strip().lower()
            if "executive summary" in h: mode = "exec"; i += 1; continue
            elif "agenda item" in h:
                mode = "items"
                if cur: data["items"].append(cur); cur = None
                i += 1; continue
            else: i += 1; continue
        if s.startswith("#### "):
            # H4 sub-heading — rendered inline as styled text within the current item
            if cur is not None:
                cur["body"].append(s)
            i += 1; continue
        if s.startswith("### "):
            if cur: data["items"].append(cur)
            h3 = s[4:].strip()
            if ":" in h3:
                n, t = h3.split(":", 1)
                cur = {"number": n.strip(), "title": t.strip(), "body": [], "next_steps": []}
            else:
                cur = {"number": h3, "title": "", "body": [], "next_steps": []}
            mode = "items"; i += 1; continue
        if not s or s == "---": i += 1; continue
        if mode == "exec":
            exec_lines.append(s)
        elif mode == "items" and cur is not None:
            ns = re.match(r"^\*\*Next Steps[:\s]*\*\*\s*(.*)", s, re.IGNORECASE)
            if ns:
                # "Next Steps" header — may have inline content or bullets on following lines
                inline = ns.group(1).strip().rstrip(".")
                if inline:
                    # Old style: **Next Steps:** item1; item2; item3
                    cur["next_steps"] = [x.strip() for x in inline.split(";") if x.strip()]
                cur["_in_next_steps"] = True
            elif cur.get("_in_next_steps") and (s.startswith("- ") or s.startswith("* ")):
                # Bullet lines following **Next Steps:**
                cur["next_steps"].append(s[2:].strip())
            else:
                cur["_in_next_steps"] = False
                cur["body"].append(s)
        i += 1
    if cur: data["items"].append(cur)
    # Preserve exec summary as list of lines (may include bullets)
    data["exec_summary"] = exec_lines
    return data


def build_and_save_briefing_v2(
    briefing_text: str,
    manifest: dict,
    meeting_folder: Path,
) -> Path:
    """
    Build a polished .docx briefing using the NEPOOL brand design system (v2).

    Same signature as build_and_save_briefing — drop-in replacement.
    Returns the Path of the saved file.
    """
    committee = manifest.get("committee", "Committee")
    dates = manifest.get("meeting_dates", [])
    date_str = _format_date_range(dates)

    # Parse markdown into structured data
    data = _v2_parse_briefing_md(briefing_text)
    data["title"] = data["title"] or committee
    data["subtitle"] = data["subtitle"] or "Meeting Briefing"
    data["date"] = data["date"] or date_str

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Twips(1008); sec.bottom_margin = Twips(1008)
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10.5)
    style.font.color.rgb = _CHARCOAL

    # Header / footer
    sec.different_first_page_header_footer = True
    # First-page: empty header, empty footer
    if sec.first_page_header.paragraphs:
        sec.first_page_header.paragraphs[0].clear()
    if sec.first_page_footer.paragraphs:
        sec.first_page_footer.paragraphs[0].clear()
    # Running header — empty (no text, just a bottom border)
    hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    hp.clear()
    # Running footer — LEFT: "Meeting Briefing" | CENTER: Page N | RIGHT: Date • Committee
    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.clear()
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_c = OxmlElement("w:tab")
    tab_c.set(qn("w:val"), "center"); tab_c.set(qn("w:pos"), str(_CONTENT_W // 2))
    tabs.append(tab_c)
    tab_r = OxmlElement("w:tab")
    tab_r.set(qn("w:val"), "right"); tab_r.set(qn("w:pos"), str(_CONTENT_W))
    tabs.append(tab_r)
    pPr.append(tabs)
    _v2_run(fp, "Meeting Briefing", size=Pt(8.5), color=_GRAY_MID)
    fp.add_run("\t")
    _v2_run(fp, "Page ", size=Pt(8.5), color=_GRAY_MID)
    pr = fp.add_run(); pr.font.name = "Calibri"; pr.font.size = Pt(8.5); pr.font.color.rgb = _GRAY_MID
    _v2_page_number(pr)
    fp.add_run("\t")
    _v2_run(fp, f"{data['date']} • {data['title']}", size=Pt(8.5), color=_GRAY_MID)
    _v2_pborder(fp, "top", 6, _CYAN_HEX, space=4)

    # Remove default empty paragraph
    if doc.paragraphs:
        doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

    # ---- Title block ----
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "top", 36, _CYAN_HEX)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_run(p, "N E P O O L", size=Pt(8), bold=True, color=_CYAN)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(4))
    _v2_run(p, data["title"], size=Pt(28), bold=True, color=_CHARCOAL)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(10)); _v2_right_tab(p)
    _v2_run(p, "Meeting Briefing", size=Pt(12), color=_CYAN)
    p.add_run("\t")
    _v2_run(p, data["date"], size=Pt(12), color=_GRAY_TEXT)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "bottom", 4, _GRAY_MID_HEX)

    # ---- Executive Summary ----
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(22), after=Pt(11))
    _v2_run(p, "EXECUTIVE SUMMARY", size=Pt(10), bold=True, color=_CHARCOAL)
    _v2_pborder(p, "bottom", 8, _CYAN_HEX, space=4)

    # Exec summary: E1 style — shaded box, each line rendered separately
    _render_v2_exec_summary(doc, data["exec_summary"])

    # ---- Agenda Item Summaries ----
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(22), after=Pt(11))
    _v2_run(p, "AGENDA ITEM SUMMARIES", size=Pt(10), bold=True, color=_CHARCOAL)
    _v2_pborder(p, "bottom", 8, _CYAN_HEX, space=4)

    for item in data["items"]:
        # Item heading
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(19), after=Pt(7))
        _v2_run(p, item["number"], size=Pt(11), bold=True, color=_CYAN)
        _v2_run(p, "  " + item["title"], size=Pt(11), bold=True, color=_CHARCOAL)
        # Body paragraphs (tables, images rendered as Word tables/pictures)
        _render_v2_body_lines(doc, item["body"])
        # Next Steps callout
        if item["next_steps"]:
            doc.add_paragraph()  # spacer
            p = doc.add_paragraph()
            _v2_pshading(p, _GRAY_BG)
            _v2_pborder(p, "top", 10, _GRAY_MID_HEX, space=6)
            _v2_pborder(p, "bottom", 10, _GRAY_MID_HEX, space=6)
            _v2_spacing(p, before=Pt(0), after=Pt(0), line=Pt(12.1))
            _v2_run(p, "NEXT STEPS", size=Pt(8.5), bold=True, color=_CYAN)
            for step in item["next_steps"]:
                p.add_run("\n")
                _v2_run(p, "–  ", size=Pt(10), color=_CHARCOAL)
                _v2_bold_runs(p, step, size=Pt(10), color=_CHARCOAL)

    # Save
    folder_name = meeting_folder.name
    filename = f"{folder_name} Briefing.docx"
    output_path = meeting_folder / filename
    doc.save(str(output_path))
    logger.info("Briefing (v2) saved: %s", output_path)
    return output_path


# ---------------------------------------------------------------------------
# On-demand generation — returns bytes for st.download_button
# ---------------------------------------------------------------------------

def generate_docx_bytes(
    briefing_text: str,
    committee: str,
    meeting_dates: list[str],
) -> bytes:
    """
    Render the v2 briefing design to a .docx and return raw bytes.
    Does not write to disk. Pass the bytes directly to st.download_button().

    Args:
        briefing_text: markdown string produced by the pipeline
        committee:     full committee name, e.g. "Markets Committee"
        meeting_dates: list of ISO date strings, e.g. ["2026-03-10"]
    """
    import io

    date_str = _format_date_range(meeting_dates)
    data = _v2_parse_briefing_md(briefing_text)
    data["title"] = data["title"] or committee
    data["subtitle"] = data["subtitle"] or "Meeting Briefing"
    data["date"] = data["date"] or date_str

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Twips(1008); sec.bottom_margin = Twips(1008)
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10.5)
    style.font.color.rgb = _CHARCOAL

    sec.different_first_page_header_footer = True
    if sec.first_page_header.paragraphs:
        sec.first_page_header.paragraphs[0].clear()
    if sec.first_page_footer.paragraphs:
        sec.first_page_footer.paragraphs[0].clear()
    # Running header — empty
    hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    hp.clear()
    # Running footer — LEFT: "Meeting Briefing" | CENTER: Page N | RIGHT: Date • Committee
    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.clear()
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_c = OxmlElement("w:tab")
    tab_c.set(qn("w:val"), "center"); tab_c.set(qn("w:pos"), str(_CONTENT_W // 2))
    tabs.append(tab_c)
    tab_r = OxmlElement("w:tab")
    tab_r.set(qn("w:val"), "right"); tab_r.set(qn("w:pos"), str(_CONTENT_W))
    tabs.append(tab_r)
    pPr.append(tabs)
    _v2_run(fp, "Meeting Briefing", size=Pt(8.5), color=_GRAY_MID)
    fp.add_run("\t")
    _v2_run(fp, "Page ", size=Pt(8.5), color=_GRAY_MID)
    pr = fp.add_run(); pr.font.name = "Calibri"; pr.font.size = Pt(8.5); pr.font.color.rgb = _GRAY_MID
    _v2_page_number(pr)
    fp.add_run("\t")
    _v2_run(fp, f"{data['date']} • {data['title']}", size=Pt(8.5), color=_GRAY_MID)
    _v2_pborder(fp, "top", 6, _CYAN_HEX, space=4)

    if doc.paragraphs:
        doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "top", 36, _CYAN_HEX)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_run(p, "N E P O O L", size=Pt(8), bold=True, color=_CYAN)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(4))
    _v2_run(p, data["title"], size=Pt(28), bold=True, color=_CHARCOAL)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(10)); _v2_right_tab(p)
    _v2_run(p, "Meeting Briefing", size=Pt(12), color=_CYAN)
    p.add_run("\t")
    _v2_run(p, data["date"], size=Pt(12), color=_GRAY_TEXT)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "bottom", 4, _GRAY_MID_HEX)

    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(22), after=Pt(11))
    _v2_run(p, "EXECUTIVE SUMMARY", size=Pt(10), bold=True, color=_CHARCOAL)
    _v2_pborder(p, "bottom", 8, _CYAN_HEX, space=4)

    _render_v2_exec_summary(doc, data["exec_summary"])

    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(22), after=Pt(11))
    _v2_run(p, "AGENDA ITEM SUMMARIES", size=Pt(10), bold=True, color=_CHARCOAL)
    _v2_pborder(p, "bottom", 8, _CYAN_HEX, space=4)

    for item in data["items"]:
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(19), after=Pt(7))
        _v2_run(p, item["number"], size=Pt(11), bold=True, color=_CYAN)
        _v2_run(p, "  " + item["title"], size=Pt(11), bold=True, color=_CHARCOAL)
        # Body paragraphs (tables, images rendered as Word tables/pictures)
        _render_v2_body_lines(doc, item["body"])
        if item["next_steps"]:
            doc.add_paragraph()
            p = doc.add_paragraph()
            _v2_pshading(p, _GRAY_BG)
            _v2_pborder(p, "top", 10, _GRAY_MID_HEX, space=6)
            _v2_pborder(p, "bottom", 10, _GRAY_MID_HEX, space=6)
            _v2_spacing(p, before=Pt(0), after=Pt(0), line=Pt(12.1))
            _v2_run(p, "NEXT STEPS", size=Pt(8.5), bold=True, color=_CYAN)
            for step in item["next_steps"]:
                p.add_run("\n")
                _v2_run(p, "–  ", size=Pt(10), color=_CHARCOAL)
                _v2_bold_runs(p, step, size=Pt(10), color=_CHARCOAL)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Deep dive / special report rendering
# ---------------------------------------------------------------------------

def _parse_deep_dive_md(text: str) -> dict:
    """
    Parse deep-dive report markdown into structured data.
    Returns {"title", "date", "sections": [{"heading", "body": [str]}]}.
    Generic section parser — not hardcoded to briefing structure.
    """
    lines = text.splitlines()
    data: dict = {"title": "", "date": "", "sections": []}
    i = 0

    # Optional H1 title
    for j, line in enumerate(lines):
        s = line.strip()
        if s.startswith("## "):
            break
        if s.startswith("# "):
            data["title"] = s[2:].strip()
            i = j + 1
            break

    # Scan for ## sections
    current_section: dict | None = None
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("## "):
            if current_section:
                data["sections"].append(current_section)
            heading = s[3:].strip()
            current_section = {"heading": heading, "body": []}
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        if current_section is not None and s:
            current_section["body"].append(s)
        elif current_section is not None and not s:
            # Preserve paragraph breaks as empty strings
            current_section["body"].append("")
        i += 1

    if current_section:
        data["sections"].append(current_section)

    return data


def generate_deep_dive_docx_bytes(
    report_md: str,
    title: str,
    document_names: list[str],
    date_range: str,
) -> bytes:
    """
    Render a deep dive report as a .docx using the NEPOOL brand design.
    Returns raw bytes for st.download_button().
    """
    import io

    data = _parse_deep_dive_md(report_md)
    data["title"] = data["title"] or title
    data["date"] = data["date"] or date_range

    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = Twips(1008); sec.bottom_margin = Twips(1008)
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(10.5)
    style.font.color.rgb = _CHARCOAL

    # Header / footer
    sec.different_first_page_header_footer = True
    if sec.first_page_header.paragraphs:
        sec.first_page_header.paragraphs[0].clear()
    if sec.first_page_footer.paragraphs:
        sec.first_page_footer.paragraphs[0].clear()
    hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    hp.clear()
    # Running footer — LEFT: "Special Report" | CENTER: Page N | RIGHT: Date • Title
    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.clear()
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab_c = OxmlElement("w:tab")
    tab_c.set(qn("w:val"), "center"); tab_c.set(qn("w:pos"), str(_CONTENT_W // 2))
    tabs.append(tab_c)
    tab_r = OxmlElement("w:tab")
    tab_r.set(qn("w:val"), "right"); tab_r.set(qn("w:pos"), str(_CONTENT_W))
    tabs.append(tab_r)
    pPr.append(tabs)
    _v2_run(fp, "Special Report", size=Pt(8.5), color=_GRAY_MID)
    fp.add_run("\t")
    _v2_run(fp, "Page ", size=Pt(8.5), color=_GRAY_MID)
    pr = fp.add_run(); pr.font.name = "Calibri"; pr.font.size = Pt(8.5); pr.font.color.rgb = _GRAY_MID
    _v2_page_number(pr)
    fp.add_run("\t")
    _v2_run(fp, f"{data['date']} • {data['title']}", size=Pt(8.5), color=_GRAY_MID)
    _v2_pborder(fp, "top", 6, _CYAN_HEX, space=4)

    # Remove default empty paragraph
    if doc.paragraphs:
        doc.paragraphs[0]._p.getparent().remove(doc.paragraphs[0]._p)

    # ---- Title block ----
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "top", 36, _CYAN_HEX)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_run(p, "N E P O O L", size=Pt(8), bold=True, color=_CYAN)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(4))
    _v2_run(p, data["title"], size=Pt(28), bold=True, color=_CHARCOAL)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(10)); _v2_right_tab(p)
    _v2_run(p, "Special Report", size=Pt(12), color=_CYAN)
    p.add_run("\t")
    _v2_run(p, data["date"], size=Pt(12), color=_GRAY_TEXT)
    p = doc.add_paragraph(); _v2_spacing(p, before=Pt(0), after=Pt(0))
    _v2_pborder(p, "bottom", 4, _GRAY_MID_HEX)

    # ---- Source documents listing ----
    if document_names:
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(10), after=Pt(6))
        _v2_run(p, "Source documents: ", size=Pt(9), bold=True, color=_CHARCOAL)
        _v2_run(p, ", ".join(document_names), size=Pt(9), color=_GRAY_TEXT)

    # ---- Sections ----
    for section in data["sections"]:
        heading = section["heading"]
        body = section["body"]

        # Section heading — cyan-underlined like EXECUTIVE SUMMARY / AGENDA ITEM SUMMARIES
        p = doc.add_paragraph(); _v2_spacing(p, before=Pt(22), after=Pt(11))
        _v2_run(p, heading.upper(), size=Pt(10), bold=True, color=_CHARCOAL)
        _v2_pborder(p, "bottom", 8, _CYAN_HEX, space=4)

        # Executive Summary gets the shaded box treatment
        if "executive summary" in heading.lower():
            _render_v2_exec_summary(doc, [l for l in body if l])
        else:
            _render_v2_body_lines(doc, body)

    # ---- Save to bytes ----
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
