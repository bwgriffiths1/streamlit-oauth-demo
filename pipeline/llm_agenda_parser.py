"""
pipeline/llm_agenda_parser.py — LLM-based agenda parsing & document matching.

Uses Claude (via tool_use) to parse agenda text into structured items and
match documents to agenda items.  Designed as a drop-in complement to the
regex parser in agenda_parser.py.

Public API:
  parse_agenda_hybrid(agenda_bytes, venue, committee, mode, config) -> (items, audit)
  llm_match_docs(agenda_items, filenames, venue, model) -> dict[str, list]
"""
import io
import json
import logging
import time as _time
from pathlib import Path

from pipeline.agenda_parser import (
    item_id_to_prefix,
    parse_agenda_from_docx,
    map_docs_to_agenda_items,
)

logger = logging.getLogger(__name__)

_PROMPTS_DIR = Path(__file__).resolve().parent.parent / "prompts"

# ---------------------------------------------------------------------------
# Tool schemas for structured output
# ---------------------------------------------------------------------------

_PARSE_TOOL = {
    "name": "parse_agenda",
    "description": "Return structured agenda items extracted from the meeting agenda text.",
    "input_schema": {
        "type": "object",
        "properties": {
            "items": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "item_id": {
                            "type": "string",
                            "description": (
                                "Dot-notation ID exactly as it appears in the agenda, "
                                "e.g. '7', '7.1', '7.1.b'.  For sub-items marked with "
                                "a letter (e.g. 'a)' or 'b)'), append the letter: '7.1.a'."
                            ),
                        },
                        "title": {
                            "type": "string",
                            "description": "Short title of the agenda item (no metadata, no parentheticals).",
                        },
                        "presenter": {
                            "type": ["string", "null"],
                            "description": "Presenter name(s) if stated, else null.",
                        },
                        "org": {
                            "type": ["string", "null"],
                            "description": "Organisation presenting (e.g. 'ISO-NE', 'Eversource'), else null.",
                        },
                        "vote_status": {
                            "type": ["string", "null"],
                            "description": "Vote tag if present (e.g. 'Future Vote', '5.0% VOTE'), else null.",
                        },
                        "wmpp_id": {
                            "type": ["string", "null"],
                            "description": "WMPP ID number if present, else null.",
                        },
                        "time_slot": {
                            "type": ["string", "null"],
                            "description": "Scheduled time range if shown, else null.",
                        },
                        "initiative_codes": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "All-caps initiative codes (e.g. 'CAR-SA', 'GISWG').",
                        },
                        "notes": {
                            "type": ["string", "null"],
                            "description": "Any remaining free-text notes after extracting other fields.",
                        },
                    },
                    "required": ["item_id", "title"],
                },
            },
        },
        "required": ["items"],
    },
}

_MATCH_TOOL = {
    "name": "match_documents",
    "description": "Assign each filename to the agenda item it belongs to.",
    "input_schema": {
        "type": "object",
        "properties": {
            "assignments": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "filename": {"type": "string"},
                        "item_id": {
                            "type": ["string", "null"],
                            "description": "The item_id this file belongs to, or null if unmatched.",
                        },
                    },
                    "required": ["filename", "item_id"],
                },
            },
        },
        "required": ["assignments"],
    },
}


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_agenda_text_docx(docx_bytes: bytes) -> str:
    """Extract table content from a .docx as a pipe-delimited markdown table."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(docx_bytes))
    parts: list[str] = []

    for idx, table in enumerate(doc.tables):
        rows_text: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " // ") for cell in row.cells]
            rows_text.append("| " + " | ".join(cells) + " |")
        if rows_text:
            parts.append(f"### Table {idx + 1}\n" + "\n".join(rows_text))

    # Also grab any paragraphs outside tables (e.g. meeting title/date header)
    header_lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            header_lines.append(text)
    header = "\n".join(header_lines[:10])  # first 10 paragraphs as context

    return f"{header}\n\n{chr(10).join(parts)}" if parts else header


def extract_agenda_text_pdf(pdf_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber with layout mode."""
    import pdfplumber

    lines: list[str] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(layout=True)
            if text:
                lines.append(text)
    return "\n\n".join(lines)


# ---------------------------------------------------------------------------
# LLM call helper (tool_use variant)
# ---------------------------------------------------------------------------

def _make_client():
    """Create an Anthropic client (reuses the same env var as summarizer)."""
    import anthropic
    return anthropic.Anthropic()


def _call_llm_tool(
    client,
    model: str,
    system_prompt: str,
    user_message: str,
    tool: dict,
    max_tokens: int = 4096,
    max_retries: int = 3,
    label: str = "",
) -> dict | None:
    """
    Call Claude with a single tool, forcing tool use.  Returns the parsed
    tool input dict, or None on failure.
    """
    for attempt in range(max_retries):
        try:
            msg = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                system=system_prompt,
                messages=[{"role": "user", "content": user_message}],
                tools=[tool],
                tool_choice={"type": "tool", "name": tool["name"]},
            )
            # Extract the tool_use block
            for block in msg.content:
                if block.type == "tool_use" and block.name == tool["name"]:
                    return block.input
            logger.warning("No tool_use block found in response%s", f" ({label})" if label else "")
            return None
        except Exception as exc:
            if "rate_limit" in str(exc).lower() and attempt < max_retries - 1:
                wait = 30 * (attempt + 1)
                logger.warning(
                    "Rate limited (attempt %d/%d), waiting %ds...%s",
                    attempt + 1, max_retries, wait, f" ({label})" if label else "",
                )
                _time.sleep(wait)
            else:
                logger.error("LLM tool call failed%s: %s", f" ({label})" if label else "", exc)
                return None
    return None


# ---------------------------------------------------------------------------
# Prompt loading
# ---------------------------------------------------------------------------

def _load_prompt(slug: str) -> str:
    path = _PROMPTS_DIR / f"{slug}.md"
    if path.exists():
        return path.read_text(encoding="utf-8")
    return ""


# ---------------------------------------------------------------------------
# LLM agenda parsing
# ---------------------------------------------------------------------------

def llm_parse_agenda(
    agenda_text: str,
    venue: str,
    committee: str,
    model: str = "claude-haiku-4-5-20251001",
) -> list[dict]:
    """
    Send raw agenda text to Claude and get back structured agenda items.

    Returns the same list[dict] contract as parse_agenda_from_docx:
      {item_id, title, prefix, auto_sub, presenter, org, vote_status,
       wmpp_id, time_slot, initiative_codes, notes, ...}
    """
    system_prompt = _load_prompt("agenda_parse_prompt")
    if not system_prompt:
        logger.error("Missing prompts/agenda_parse_prompt.md — cannot run LLM parsing")
        return []

    user_msg = (
        f"Venue: {venue}\nCommittee: {committee}\n\n"
        f"--- AGENDA TEXT ---\n{agenda_text}"
    )

    client = _make_client()
    result = _call_llm_tool(
        client, model, system_prompt, user_msg, _PARSE_TOOL,
        max_tokens=8192, label=f"parse-{venue}-{committee}",
    )
    if not result or "items" not in result:
        return []

    # Post-process: normalise IDs to match regex conventions, compute prefix
    items: list[dict] = []
    for raw in result["items"]:
        item_id = raw.get("item_id", "").strip().rstrip("*")
        if not item_id:
            continue
        # Strip trailing ".0" — regex convention (2.0 → 2, but 2.0.1 kept)
        parts = item_id.split(".")
        if len(parts) == 2 and parts[1] == "0":
            item_id = parts[0]
        # Strip leading committee letter prefix (e.g. "A-7" → "7")
        import re as _re
        item_id = _re.sub(r"^[A-Za-z]-", "", item_id)

        title = raw.get("title", "").strip()
        if not title:
            continue

        prefix = item_id_to_prefix(item_id)
        items.append({
            "item_id": item_id,
            "title": title,
            "prefix": prefix,
            "auto_sub": False,
            "presenter": raw.get("presenter"),
            "org": raw.get("org"),
            "vote_status": raw.get("vote_status"),
            "wmpp_id": raw.get("wmpp_id"),
            "time_slot": raw.get("time_slot"),
            "initiative_codes": raw.get("initiative_codes", []),
            "other_tags": [],
            "notes": raw.get("notes"),
            "meeting_number": None,
        })

    logger.info("LLM parsed %d agenda item(s) for %s/%s", len(items), venue, committee)
    return items


# ---------------------------------------------------------------------------
# LLM document matching
# ---------------------------------------------------------------------------

def llm_match_docs(
    agenda_items: list[dict],
    filenames: list[str],
    venue: str,
    model: str = "claude-haiku-4-5-20251001",
) -> dict[str, list]:
    """
    Use Claude to assign filenames to agenda items.

    Returns the same bucket dict as map_docs_to_agenda_items:
      {prefix: [{"filename": ...}, ...], "other": [...]}
    """
    system_prompt = _load_prompt("doc_match_prompt")
    if not system_prompt:
        logger.error("Missing prompts/doc_match_prompt.md — cannot run LLM matching")
        return {"other": [{"filename": f} for f in filenames]}

    # Build context for the LLM
    items_desc = "\n".join(
        f"  - item_id={it['item_id']}  prefix={it.get('prefix', '?')}  title={it['title']}"
        for it in agenda_items
    )
    files_desc = "\n".join(f"  - {f}" for f in filenames)

    user_msg = (
        f"Venue: {venue}\n\n"
        f"AGENDA ITEMS:\n{items_desc}\n\n"
        f"FILENAMES:\n{files_desc}"
    )

    client = _make_client()
    result = _call_llm_tool(
        client, model, system_prompt, user_msg, _MATCH_TOOL,
        max_tokens=4096, label=f"match-{venue}",
    )
    if not result or "assignments" not in result:
        return {"other": [{"filename": f} for f in filenames]}

    # Build item_id → prefix lookup (normalise IDs for resilient matching)
    def _norm(iid: str) -> str:
        n = iid.lower().rstrip("*").strip()
        parts = n.split(".")
        if len(parts) == 2 and parts[1] == "0":
            n = parts[0]
        return n

    id_to_prefix = {_norm(it["item_id"]): it.get("prefix") for it in agenda_items}

    buckets: dict[str, list] = {"other": []}
    for item in agenda_items:
        p = item.get("prefix")
        if p:
            buckets[p] = []

    for assignment in result["assignments"]:
        fn = assignment.get("filename", "")
        raw_id = assignment.get("item_id")
        prefix = id_to_prefix.get(_norm(raw_id)) if raw_id else None
        if prefix and prefix in buckets:
            buckets[prefix].append({"filename": fn})
        else:
            buckets["other"].append({"filename": fn})

    return {k: v for k, v in buckets.items() if v}


# ---------------------------------------------------------------------------
# Reconciliation
# ---------------------------------------------------------------------------

def reconcile_results(
    regex_items: list[dict],
    llm_items: list[dict],
) -> dict:
    """
    Compare regex and LLM parse results.

    Returns:
      {
        "regex_count": int,
        "llm_count": int,
        "matched": [{"item_id": ..., "diffs": {field: (regex_val, llm_val)}}],
        "regex_only": [item_id, ...],   # in regex but not LLM
        "llm_only": [item_id, ...],     # in LLM but not regex
        "agreement_pct": float,          # % of items that fully agree
      }
    """
    # Normalise item_ids for matching: lowercase, strip trailing .0
    def _norm_id(iid: str) -> str:
        n = iid.lower().rstrip("*").strip()
        parts = n.split(".")
        if len(parts) == 2 and parts[1] == "0":
            n = parts[0]
        return n

    regex_by_id = {_norm_id(it["item_id"]): it for it in regex_items}
    llm_by_id = {_norm_id(it["item_id"]): it for it in llm_items}

    all_ids = set(regex_by_id) | set(llm_by_id)
    regex_only = sorted(set(regex_by_id) - set(llm_by_id))
    llm_only = sorted(set(llm_by_id) - set(regex_by_id))

    compare_fields = ["title", "presenter", "org", "vote_status", "wmpp_id"]
    # Title comparison is case-insensitive (regex preserves source casing,
    # LLM may normalise to title case).
    case_insensitive_fields = {"title"}
    matched = []
    full_agree = 0

    for iid in sorted(set(regex_by_id) & set(llm_by_id)):
        r = regex_by_id[iid]
        l = llm_by_id[iid]
        diffs = {}
        for field in compare_fields:
            rv = (r.get(field) or "").strip() if r.get(field) else None
            lv = (l.get(field) or "").strip() if l.get(field) else None
            if field in case_insensitive_fields:
                if (rv or "").lower() != (lv or "").lower():
                    diffs[field] = (rv, lv)
            elif rv != lv:
                diffs[field] = (rv, lv)
        matched.append({"item_id": iid, "diffs": diffs})
        if not diffs:
            full_agree += 1

    total = len(all_ids) or 1
    return {
        "regex_count": len(regex_items),
        "llm_count": len(llm_items),
        "matched": matched,
        "regex_only": regex_only,
        "llm_only": llm_only,
        "agreement_pct": round(full_agree / total * 100, 1),
    }


def _merge_results(
    regex_items: list[dict],
    llm_items: list[dict],
    reconciliation: dict,
) -> list[dict]:
    """
    Merge regex and LLM results with these rules:
      - item_id and prefix: always from regex (deterministic)
      - metadata fields (presenter, org, vote_status, wmpp_id): prefer LLM
      - items found only by LLM: include them (regex missed)
      - items found only by regex: keep them (LLM missed)
      - title: prefer LLM (often cleaner)
    """
    def _norm_id(iid: str) -> str:
        n = iid.lower().rstrip("*").strip()
        parts = n.split(".")
        if len(parts) == 2 and parts[1] == "0":
            n = parts[0]
        return n

    regex_by_id = {_norm_id(it["item_id"]): it for it in regex_items}
    llm_by_id = {_norm_id(it["item_id"]): it for it in llm_items}

    merged: list[dict] = []
    seen_ids: set[str] = set()

    # Process in regex order first (preserves sequence)
    for item in regex_items:
        iid = _norm_id(item["item_id"])
        seen_ids.add(iid)
        if iid in llm_by_id:
            llm_item = llm_by_id[iid]
            merged.append({
                **item,
                # Prefer LLM for metadata fields
                "title": llm_item.get("title") or item["title"],
                "presenter": llm_item.get("presenter") or item.get("presenter"),
                "org": llm_item.get("org") or item.get("org"),
                "vote_status": llm_item.get("vote_status") or item.get("vote_status"),
                "wmpp_id": llm_item.get("wmpp_id") or item.get("wmpp_id"),
                "initiative_codes": llm_item.get("initiative_codes") or item.get("initiative_codes", []),
                "notes": llm_item.get("notes") or item.get("notes"),
            })
        else:
            merged.append(item)

    # Add items found only by LLM (regex missed them)
    for raw_iid in reconciliation.get("llm_only", []):
        iid = _norm_id(raw_iid)
        if iid not in seen_ids and iid in llm_by_id:
            llm_item = {**llm_by_id[iid]}
            # Compute prefix server-side
            llm_item["prefix"] = item_id_to_prefix(raw_iid)
            merged.append(llm_item)

    return merged


# ---------------------------------------------------------------------------
# Hybrid entry point
# ---------------------------------------------------------------------------

def parse_agenda_hybrid(
    agenda_bytes: bytes,
    venue: str,
    committee: str,
    mode: str = "llm_verify",
    config: dict | None = None,
) -> tuple[list[dict], dict | None]:
    """
    Parse an agenda using the configured mode.

    Args:
        agenda_bytes: raw .docx (or .pdf) bytes
        venue: e.g. "ISO-NE", "NYISO"
        committee: e.g. "MC", "RC"
        mode: "regex_only", "llm_only", "llm_verify", "llm_fallback"
        config: optional config dict with agenda_parsing section

    Returns:
        (parsed_items, audit_dict_or_None)
    """
    cfg = (config or {}).get("agenda_parsing", {})
    parse_model = cfg.get("parse_model", "claude-haiku-4-5-20251001")
    escalation_model = cfg.get("escalation_model", "claude-sonnet-4-6")

    # ── Detect format from magic bytes ─────────────────────────────────────
    is_pdf = agenda_bytes[:5] == b"%PDF-"

    # ── regex_only ──────────────────────────────────────────────────────────
    if mode == "regex_only":
        if is_pdf:
            logger.warning("PDF agenda in regex_only mode — switching to llm_only")
            mode = "llm_only"
        else:
            items = parse_agenda_from_docx(agenda_bytes)
            return items, None

    # ── Extract text for LLM ────────────────────────────────────────────────
    if is_pdf:
        agenda_text = extract_agenda_text_pdf(agenda_bytes)
    else:
        agenda_text = extract_agenda_text_docx(agenda_bytes)

    if not agenda_text.strip():
        logger.warning("No text extracted from agenda — falling back to regex")
        return parse_agenda_from_docx(agenda_bytes), None

    # ── llm_only ────────────────────────────────────────────────────────────
    if mode == "llm_only":
        items = llm_parse_agenda(agenda_text, venue, committee, model=parse_model)
        if not items:
            logger.warning("LLM parse returned no items — falling back to regex")
            items = parse_agenda_from_docx(agenda_bytes)
            return items, {"mode": "llm_only", "fallback": "regex", "reason": "empty_llm_result"}
        return items, {"mode": "llm_only", "llm_count": len(items)}

    # ── llm_verify ──────────────────────────────────────────────────────────
    if mode == "llm_verify":
        # Step 1: regex pass (skip for PDF — regex can't parse PDFs)
        if is_pdf:
            logger.info("PDF agenda — running LLM-only (no regex available)")
            items = llm_parse_agenda(agenda_text, venue, committee, model=parse_model)
            if not items:
                logger.warning("LLM parse returned no items for PDF agenda")
            return items, {"mode": "llm_only", "reason": "pdf_agenda", "llm_count": len(items)}
        regex_items = parse_agenda_from_docx(agenda_bytes)

        # Step 2: LLM pass
        llm_items = llm_parse_agenda(agenda_text, venue, committee, model=parse_model)

        if not llm_items:
            logger.warning("LLM parse failed in verify mode — using regex results")
            return regex_items, {"mode": "llm_verify", "fallback": "regex_only", "reason": "llm_failed"}

        # Step 3: reconcile
        audit = reconcile_results(regex_items, llm_items)
        audit["mode"] = "llm_verify"
        audit["model"] = parse_model

        # Step 4: if significant structural disagreement, escalate to stronger model
        # Only count missing/extra items as structural — field-level diffs
        # (presenter, vote_status) are typically LLM doing better, not errors.
        structural_disagreements = len(audit["regex_only"]) + len(audit["llm_only"])

        if structural_disagreements > 2:
            logger.info(
                "Significant disagreement (%d issues) — escalating to %s",
                structural_disagreements, escalation_model,
            )
            escalated_items = llm_parse_agenda(
                agenda_text, venue, committee, model=escalation_model,
            )
            if escalated_items:
                llm_items = escalated_items
                audit["escalated"] = True
                audit["escalation_model"] = escalation_model
                # Re-reconcile with escalated results
                escalated_audit = reconcile_results(regex_items, llm_items)
                audit.update({
                    "escalated_matched": escalated_audit["matched"],
                    "escalated_regex_only": escalated_audit["regex_only"],
                    "escalated_llm_only": escalated_audit["llm_only"],
                    "escalated_agreement_pct": escalated_audit["agreement_pct"],
                })

        # Step 5: merge
        merged = _merge_results(regex_items, llm_items, audit)
        audit["final_count"] = len(merged)
        return merged, audit

    # ── llm_fallback ────────────────────────────────────────────────────────
    if mode == "llm_fallback":
        if is_pdf:
            regex_items = []
        else:
            try:
                regex_items = parse_agenda_from_docx(agenda_bytes)
            except Exception as exc:
                logger.warning("Regex parse failed: %s — trying LLM", exc)
                regex_items = []

        # Trigger LLM if regex failed or looks suspicious (very few items)
        if len(regex_items) < 3:
            logger.info(
                "Regex produced only %d items — running LLM fallback", len(regex_items),
            )
            llm_items = llm_parse_agenda(agenda_text, venue, committee, model=parse_model)
            if llm_items and len(llm_items) > len(regex_items):
                audit = reconcile_results(regex_items, llm_items)
                audit["mode"] = "llm_fallback"
                audit["reason"] = "regex_insufficient"
                merged = _merge_results(regex_items, llm_items, audit)
                return merged, audit
            # LLM didn't help — use whatever regex gave us
            return regex_items, {"mode": "llm_fallback", "reason": "llm_no_improvement"}

        return regex_items, None

    logger.error("Unknown agenda_parsing mode: %s — using regex_only", mode)
    return parse_agenda_from_docx(agenda_bytes), None
