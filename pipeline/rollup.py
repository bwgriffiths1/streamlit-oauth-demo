"""
rollup.py — Phase 2: agenda-item and meeting-level summarization.

Layer 2: groups document summaries by agenda item prefix and calls Claude
         to produce a per-item synthesis. Results stored in Postgres.
Layer 3: sends all item summaries to Claude to produce the meeting-level
         briefing text. Stored in Postgres; .docx generated on-demand from UI.
"""
import logging
import re
from pathlib import Path

import anthropic
import requests as _requests
from docx import Document as DocxDocument

import pipeline.db as db
from pipeline.downloader import download_file_temp
from pipeline.utils import fix_markdown

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Agenda parsing (unchanged — operates on a temp-downloaded .docx file path)
# ---------------------------------------------------------------------------

def _item_id_to_prefix(raw_id: str) -> str | None:
    raw = raw_id.rstrip("*").strip()
    raw = re.sub(r"^[A-Za-z]-", "", raw)
    if not raw or not raw[0].isdigit():
        return None
    parts = raw.split(".")
    if len(parts) == 2 and parts[1] == "0":
        parts = [parts[0]]
    if len(parts) > 1 and re.fullmatch(r"[A-Za-z]+", parts[-1]):
        parts = parts[:-1]
    prefix = "a" + parts[0].zfill(2)
    if len(parts) > 1:
        prefix += "." + ".".join(parts[1:])
    return prefix


_ITEM_ID_RE = re.compile(r"^(\d+(?:\.\d+)*(?:\.[A-Za-z]+)?)\*?\s*$")

_PARA_PATTERNS = [
    re.compile(r"^([A-Z]-\d+(?:\.\d+)*)\s*[.:\-]\s+(.+)"),
    re.compile(r"^(\d+(?:\.\d+)*)\s*[.:\-]\s+(.+)"),
    re.compile(r"^([A-Z])\s*\.\s+(.+)"),
]


def _parse_agenda_from_tables(doc: DocxDocument) -> list[dict]:
    seen: set[tuple[str, str]] = set()
    items: list[dict] = []
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue
            id_cell = cells[0].text.strip()
            if not id_cell:
                continue
            if not _ITEM_ID_RE.match(id_cell):
                continue
            raw_id = id_cell.rstrip("*").strip()
            desc_cell = cells[1].text.strip() if len(cells) > 1 else id_cell
            title = desc_cell.splitlines()[0].strip().rstrip(".;:*")
            if not title:
                continue
            key = (raw_id, title)
            if key in seen:
                continue
            seen.add(key)
            prefix = _item_id_to_prefix(id_cell)
            items.append({"item_id": raw_id, "title": title, "prefix": prefix})
    return items


def parse_agenda(agenda_path: Path) -> list[dict]:
    """
    Parse a NEPOOL agenda .docx to extract an ordered list of agenda items.
    agenda_path is a temp file path — no permanent storage assumed.
    """
    doc = DocxDocument(str(agenda_path))
    items = _parse_agenda_from_tables(doc)
    if items:
        logger.info("parse_agenda: %d items from tables in %s", len(items), agenda_path.name)
        return items

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        for pat in _PARA_PATTERNS:
            m = pat.match(text)
            if m:
                raw_id = m.group(1)
                title = m.group(2).strip().rstrip(".;:")
                items.append({
                    "item_id": raw_id,
                    "title": title,
                    "prefix": _item_id_to_prefix(raw_id),
                })
                break

    if items:
        logger.info("parse_agenda: %d items from paragraphs in %s", len(items), agenda_path.name)
        return items

    logger.warning("parse_agenda: no items found in %s — using fallback", agenda_path.name)
    return [{"item_id": "0", "title": "Full Meeting", "prefix": None}]


# ---------------------------------------------------------------------------
# Mapping documents to agenda items (now uses DB summary dicts, not disk files)
# ---------------------------------------------------------------------------

def map_docs_to_agenda_items(
    doc_summary_rows: list[dict],
    agenda_items: list[dict],
) -> dict[str, list[dict]]:
    """
    Group document summary DB rows under their agenda item by filename prefix.

    doc_summary_rows: rows from get_all_document_summaries_for_meeting()
                      — each has 'filename' and 'summary_text'.

    Returns dict: {item_id → [row_dict, ...]}.
    Unmatched docs accumulate under the key "other".
    """
    prefix_map: dict[str, str] = {}
    for item in agenda_items:
        if item.get("prefix"):
            existing = prefix_map.get(item["prefix"])
            if existing is None or item["item_id"].count(".") < existing.count("."):
                prefix_map[item["prefix"]] = item["item_id"]

    grouped: dict[str, list] = {item["item_id"]: [] for item in agenda_items}
    grouped.setdefault("other", [])

    for row in doc_summary_rows:
        filename_lower = row["filename"].lower()
        matched_item_id = None
        matched_len = 0
        for prefix, item_id in prefix_map.items():
            if filename_lower.startswith(prefix) and len(prefix) > matched_len:
                matched_item_id = item_id
                matched_len = len(prefix)

        if matched_item_id:
            grouped[matched_item_id].append(row)
        else:
            grouped["other"].append(row)

    return {k: v for k, v in grouped.items() if v}


# ---------------------------------------------------------------------------
# Agenda-item level summarization
# ---------------------------------------------------------------------------

def summarize_agenda_item(
    item: dict,
    doc_entries: list[dict],
    client: anthropic.Anthropic,
    model: str,
    agenda_item_prompt: str,
    max_tokens: int = 1024,
) -> str:
    doc_summaries_text = "\n\n---\n\n".join(
        f"**{d['filename']}**\n\n{d['summary_text']}" for d in doc_entries
    )
    prompt = agenda_item_prompt.format(
        item_id=item["item_id"],
        title=item["title"],
        doc_summaries=doc_summaries_text,
    )
    text = ""
    with client.messages.stream(
        model=model,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}],
    ) as stream:
        for text_chunk in stream.text_stream:
            text += text_chunk
    return fix_markdown(text)


# ---------------------------------------------------------------------------
# Meeting-level briefing text
# ---------------------------------------------------------------------------

def generate_meeting_briefing_text(
    item_results: dict,
    briefing_prompt: str,
    client: anthropic.Anthropic,
    model: str,
    max_tokens: int = 4096,
) -> str:
    def _sort_key(item_id: str) -> tuple:
        if item_id == "other":
            return (999,)
        parts = re.split(r"[.\-]", re.sub(r"^[A-Za-z]-", "", item_id))
        result = []
        for p in parts:
            try:
                result.append(int(p))
            except ValueError:
                result.append(p)
        return tuple(result)

    items_block_parts = []
    for item_id in sorted(item_results.keys(), key=_sort_key):
        result = item_results[item_id]
        title = result["item"].get("title", item_id)
        items_block_parts.append(
            f"## Item {item_id}: {title}\n\n{result['summary_text']}"
        )
    items_block = "\n\n---\n\n".join(items_block_parts)

    if "[AGENDA ITEMS]" not in briefing_prompt:
        logger.warning("Briefing prompt has no [AGENDA ITEMS] placeholder — appending.")
        full_prompt = briefing_prompt + "\n\n" + items_block
    else:
        full_prompt = briefing_prompt.replace("[AGENDA ITEMS]", items_block)

    text = ""
    with client.messages.stream(
        model=model,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": full_prompt}],
    ) as stream:
        for text_chunk in stream.text_stream:
            text += text_chunk
    return fix_markdown(text)


# ---------------------------------------------------------------------------
# Pipeline entry point
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Committee → prompt slug helpers
# ---------------------------------------------------------------------------

_COMM_SHORT = {
    "Markets Committee":      "mc",
    "Participants Committee": "npc",
    "Reliability Committee":  "rc",
}


def _get_committee_prompts(committee: str) -> tuple[str, str]:
    """
    Return (briefing_prompt, agenda_item_prompt) for the given committee.
    Falls back to mc_* slugs if committee-specific prompt not found.
    Falls back to legacy slug 'mc_briefing_prompt' / 'agenda_item_prompt'.
    Prepends general_context_prompt if it is non-empty.
    """
    short = _COMM_SHORT.get(committee, "mc")

    def _load(slug: str, fallback: str) -> str:
        row = db.get_prompt_by_slug(slug) or db.get_prompt_by_slug(fallback)
        return (row["content"] if row else "") or ""

    briefing_prompt     = _load(f"{short}_briefing_prompt",     "mc_briefing_prompt")
    agenda_item_prompt  = _load(f"{short}_agenda_item_prompt",  "agenda_item_prompt")

    # Prepend general context if non-empty
    ctx_row = db.get_prompt_by_slug("general_context_prompt")
    general_ctx = (ctx_row["content"] if ctx_row else "").strip()
    if general_ctx:
        briefing_prompt    = general_ctx + "\n\n" + briefing_prompt
        agenda_item_prompt = general_ctx + "\n\n" + agenda_item_prompt

    return briefing_prompt, agenda_item_prompt


# ---------------------------------------------------------------------------
# Keyword extraction (Layer 4 — cheap post-processing step)
# ---------------------------------------------------------------------------

def _extract_and_store_keywords(
    meeting_id: int,
    briefing_text: str,
    client: anthropic.Anthropic,
    model: str = "claude-haiku-4-5",
) -> None:
    """Call Claude to extract ISO-NE keywords from the briefing; store in DB."""
    import json

    kw_row = db.get_prompt_by_slug("keyword_extraction_prompt")
    if not kw_row:
        logger.warning("keyword_extraction_prompt not found in DB — skipping")
        return

    prompt = kw_row["content"].replace("{briefing_text}", briefing_text)
    try:
        msg = client.messages.create(
            model=model,
            max_tokens=512,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        # Strip markdown code fences if present
        raw = raw.strip("`").strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
        keywords = json.loads(raw)
        if isinstance(keywords, list):
            keywords = [str(k) for k in keywords if k]
            db.set_ai_keywords(meeting_id, keywords)
            logger.info("  extracted %d keywords for meeting %d", len(keywords), meeting_id)
    except Exception as exc:
        logger.warning("  keyword extraction failed: %s", exc)


# ---------------------------------------------------------------------------
# Re-rollup helpers (for use from UI after editing item summaries)
# ---------------------------------------------------------------------------

def regenerate_briefing_from_existing_items(
    meeting_id: int,
    committee: str,
    client: anthropic.Anthropic,
    config: dict,
) -> None:
    """
    Regenerate the meeting-level briefing using existing item summaries (Layer 3 only).
    Prefers user_text over summary_text for each item if edits exist.
    Stores the new briefing as the AI text (clears any prior user_text edit).
    """
    summ_cfg = config.get("summarization", {})
    models   = summ_cfg.get("models", {})
    briefing_model = models.get("briefing", "claude-opus-4-5")

    briefing_prompt, _ = _get_committee_prompts(committee)

    rows = db.get_all_agenda_item_summaries_for_meeting(meeting_id)
    if not rows:
        logger.warning("regenerate_briefing: no item summaries found for meeting %d", meeting_id)
        return

    item_results = {}
    for row in rows:
        item_id = row["item_number"].replace("_", ".")
        text    = row.get("user_text") or row["summary_text"]
        item_results[item_id] = {
            "item": {"item_id": item_id, "title": row.get("item_title", item_id)},
            "summary_text": text,
        }

    briefing_text = generate_meeting_briefing_text(
        item_results=item_results,
        briefing_prompt=briefing_prompt,
        client=client,
        model=briefing_model,
    )
    db.upsert_briefing(meeting_id, briefing_text, model_used=briefing_model)
    # Clear any prior user edit so the new AI text is displayed
    db.update_briefing_user_text(meeting_id, None)
    logger.info("  briefing regenerated from existing items for meeting %d", meeting_id)
    _extract_and_store_keywords(meeting_id, briefing_text, client, model=models.get("rollup", "claude-haiku-4-5"))


def regenerate_item_from_existing_docs(
    meeting_id: int,
    item_number: str,
    item_title: str,
    doc_summary_rows: list[dict],
    committee: str,
    client: anthropic.Anthropic,
    config: dict,
) -> None:
    """
    Regenerate one agenda-item summary from existing doc summaries (Layer 2 only).
    Prefers user_text over summary_text for each doc.
    """
    summ_cfg    = config.get("summarization", {})
    rollup_model = summ_cfg.get("models", {}).get("rollup", "claude-haiku-4-5")
    max_tokens  = summ_cfg.get("max_response_tokens", 1024)

    _, agenda_item_prompt = _get_committee_prompts(committee)

    # Prefer user_text if present
    effective_docs = [
        {**r, "summary_text": r.get("user_text") or r["summary_text"]}
        for r in doc_summary_rows
    ]

    item_id_display = item_number.replace("_", ".")
    item = {"item_id": item_id_display, "title": item_title}
    summary_text = summarize_agenda_item(
        item=item,
        doc_entries=effective_docs,
        client=client,
        model=rollup_model,
        agenda_item_prompt=agenda_item_prompt,
        max_tokens=max_tokens,
    )
    db.upsert_agenda_item_summary(
        meeting_id=meeting_id,
        item_number=item_number,
        summary_text=summary_text,
        item_title=item_title,
        model_used=rollup_model,
    )
    # Clear any prior user edit
    db.update_item_summary_user_text(meeting_id, item_number, None)
    logger.info("  item %s regenerated from existing docs for meeting %d", item_number, meeting_id)


def run_rollup(
    meeting_id: int,
    client: anthropic.Anthropic,
    config: dict,
    briefing_prompt: str,
    agenda_item_prompt: str,
    session: _requests.Session | None = None,
    committee: str = "Markets Committee",
) -> None:
    """
    Full Layer 2 + Layer 3 pipeline for one meeting.

    1. Find the agenda .docx in documents, download to temp, parse it.
    2. Map completed doc summaries (from DB) to agenda items by filename prefix.
    3. For each non-empty item: call Claude → upsert to agenda_item_summaries.
    4. Generate meeting-level briefing → upsert to briefings.
    5. Extract keywords from briefing → upsert to meeting_keywords.
    6. Mark meeting summary_status = 'complete'.

    briefing_prompt and agenda_item_prompt are loaded from DB by the caller;
    if the caller passes empty strings the committee-specific prompts are used.
    """
    if session is None:
        session = _requests.Session()

    summ_cfg = config.get("summarization", {})
    models = summ_cfg.get("models", {})
    rollup_model = models.get("rollup", "claude-haiku-4-5")
    briefing_model = models.get("briefing", "claude-opus-4-5")
    max_response_tokens = summ_cfg.get("max_response_tokens", 4096)

    # Use committee-specific prompts (caller may pass empty strings as signal to auto-resolve)
    if not briefing_prompt or not agenda_item_prompt:
        briefing_prompt, agenda_item_prompt = _get_committee_prompts(committee)

    # --- Step 1: parse agenda ---
    docs = db.get_documents_for_meeting(meeting_id)
    agenda_doc = next(
        (d for d in docs
         if "agenda" in d["filename"].lower()
         and d["filename"].lower().endswith(".docx")),
        None,
    )

    if agenda_doc and agenda_doc.get("source_url"):
        try:
            with download_file_temp(
                url=agenda_doc["source_url"],
                filename=agenda_doc["filename"],
                referer_url=agenda_doc["source_url"],
                session=session,
            ) as tmp_path:
                if tmp_path:
                    agenda_items = parse_agenda(Path(tmp_path))
                else:
                    agenda_items = [{"item_id": "0", "title": "Full Meeting", "prefix": None}]
        except Exception as exc:
            logger.warning("run_rollup: agenda parse failed (%s) — using fallback", exc)
            agenda_items = [{"item_id": "0", "title": "Full Meeting", "prefix": None}]
    else:
        logger.warning("run_rollup: no agenda .docx found — using fallback")
        agenda_items = [{"item_id": "0", "title": "Full Meeting", "prefix": None}]

    # --- Step 2: map doc summaries to agenda items ---
    doc_summaries = db.get_all_document_summaries_for_meeting(meeting_id)
    grouped = map_docs_to_agenda_items(doc_summaries, agenda_items)

    if not grouped:
        logger.warning("run_rollup: no summarized documents to roll up")
        db.set_meeting_summary_status(meeting_id, "complete")
        return

    item_lookup = {item["item_id"]: item for item in agenda_items}

    # --- Step 3: per-item summaries ---
    # Clear stale summaries from any previous run before writing new ones.
    db.clear_agenda_item_summaries(meeting_id)
    item_results: dict[str, dict] = {}
    for item_id, doc_entries in grouped.items():
        item = item_lookup.get(item_id, {"item_id": item_id, "title": item_id})
        logger.info(
            "  item %s (%s): rolling up %d doc(s)",
            item_id, item["title"], len(doc_entries),
        )
        try:
            summary_text = summarize_agenda_item(
                item=item,
                doc_entries=doc_entries,
                client=client,
                model=rollup_model,
                agenda_item_prompt=agenda_item_prompt,
                max_tokens=max_response_tokens,
            )
            # safe_id: "2.0" → "2_0", used as item_number in DB
            safe_id = item_id.replace(".", "_").replace("-", "_")
            db.upsert_agenda_item_summary(
                meeting_id=meeting_id,
                item_number=safe_id,
                summary_text=summary_text,
                item_title=item["title"],
                model_used=rollup_model,
            )
            item_results[item_id] = {
                "item": item,
                "summary_text": summary_text,
            }
            logger.info("  saved to DB: item %s", item_id)
        except anthropic.APIError as exc:
            logger.error("  API error on item %s: %s", item_id, exc)
        except Exception as exc:
            logger.error("  error on item %s: %s", item_id, exc)

    # --- Step 4: meeting-level briefing ---
    if not item_results:
        logger.warning("run_rollup: no item summaries produced; skipping briefing")
    else:
        logger.info("  generating meeting briefing...")
        try:
            briefing_text = generate_meeting_briefing_text(
                item_results=item_results,
                briefing_prompt=briefing_prompt,
                client=client,
                model=briefing_model,
                max_tokens=4096,
            )
            db.upsert_briefing(meeting_id, briefing_text, model_used=briefing_model)
            logger.info("  briefing saved to DB.")

            # --- Step 5: keyword extraction ---
            try:
                _extract_and_store_keywords(meeting_id, briefing_text, client, model=rollup_model)
            except Exception as exc:
                logger.warning("  keyword extraction error (non-fatal): %s", exc)

        except Exception as exc:
            logger.error("  error generating briefing: %s", exc)

    db.set_meeting_summary_status(meeting_id, "complete")
