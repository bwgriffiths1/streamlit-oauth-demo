"""
pipeline/agenda_parser.py

Shared agenda parsing utilities.  Used by:
  - main.py  (Phase 0.5 — store skeleton in DB)
  - pages/08_agenda_debug.py  (diagnostic display)

Public API:
  parse_agenda_from_docx(path_or_bytes) -> list[dict]
  map_docs_to_agenda_items(doc_rows, agenda_items) -> dict[str, list]
"""
import re
from pathlib import Path

from docx import Document as DocxDocument

# ---------------------------------------------------------------------------
# Prefix helpers
# ---------------------------------------------------------------------------

def item_id_to_prefix(raw_id: str) -> str | None:
    """
    Convert an agenda item ID (e.g. "7.1.b") to its filename prefix ("a07.1.b").

    Rules:
      - Strip trailing *
      - Strip leading committee letter prefix like "A-"
      - Zero-pad top-level number to 2 digits
      - Preserve all sub-parts (dots and letters) — never strip terminal letters
      - Strip trailing ".0" only when it is the sole sub-part (e.g. "2.0" → "a02")
    """
    raw = raw_id.rstrip("*").strip()
    raw = re.sub(r"^[A-Za-z]-", "", raw)
    if not raw or not raw[0].isdigit():
        return None
    parts = raw.split(".")
    if len(parts) == 2 and parts[1] == "0":
        parts = [parts[0]]
    prefix = "a" + parts[0].zfill(2)
    if len(parts) > 1:
        prefix += "." + ".".join(p.lower() for p in parts[1:])
    return prefix


# Matches agenda ID cells: "7", "7.1", "7.1.b", "1.A", "4.1*", "2.0"
_ITEM_ID_RE = re.compile(r"^(\d+(?:\.\d+)*(?:\.[A-Za-z])?)\*?\s*$")

# Sub-letter at start of description cell: "a) Title text" or "b) Title"
_DESC_SUB_LETTER_RE = re.compile(r"^([a-zA-Z])\)\s*(.*)", re.DOTALL)

# ---------------------------------------------------------------------------
# Metadata extraction regexes
# ---------------------------------------------------------------------------

# One parenthetical group — handles one level of nesting e.g. "(FERC Order (EL25-106))"
_PAREN_RE = re.compile(r"\(([^()]*(?:\([^()]*\)[^()]*)*)\)")

# Presenter patterns:
#   "(ISO Staff)"                           → org=ISO, presenter=ISO Staff
#   "(ISO: Name)"  "(ISO-NE: Name1, Name2)" → org=ISO/ISO-NE, presenter=names
#   "(Rhode Island Energy: Boris Y.)"       → org=Rhode Island Energy, presenter=Boris Y.
#   "(GIS Working Group Chair: Jamie S.)"   → org=GIS Working Group Chair, presenter=Jamie S.
# Two forms:
#   1. "ISO[-NE] Staff" with no colon
#   2. "OrgName: PersonNames" for any org (2–60 chars before colon, no parens)
_ISO_STAFF_RE = re.compile(r"^ISO(?:-NE)?\s+Staff$", re.IGNORECASE)
_ORG_PRESENTER_RE = re.compile(r"^([^:()]{2,60}):\s*(.+)$")

# Vote / status tags
_VOTE_RE = re.compile(
    r"^(Future Vote|[\d.]+\s*%\s*VOTE?|NEPOOL Vote|Notification)$", re.IGNORECASE
)

# Meeting ordinal: "1st MC Mtg", "10th MC/RC Mtg", "3rd MC/RC Meeting"
_MTG_NUM_RE = re.compile(
    r"^(\d+(?:st|nd|rd|th)\s+(?:MC|RC)(?:/(?:RC|MC))?\s+(?:Mtg\.?|Meeting))$",
    re.IGNORECASE,
)

# WMPP ID cross-reference
_WMPP_RE = re.compile(r"^WMPP\s+ID:\s*(\d+)$", re.IGNORECASE)

# All-caps initiative code with separator: "(CAR-SA)", "(CAR/SA)", "(GISWG)"
_INIT_CODE_RE = re.compile(r"^[A-Z]{2,}(?:[-/][A-Z]{2,})+$")


def _has_metadata_parens(text: str) -> bool:
    """Return True if text has at least one parenthetical matching a known metadata pattern.

    Used to detect empty-ID sub-item rows (e.g. RC agendas where a sub-topic appears
    on a separate row with no ID cell but carries a presenter / meeting number).
    Avoids false-positives on Roman-numeral lists like '(i)', '(ii)'.
    """
    for m in _PAREN_RE.finditer(text):
        content = m.group(1).strip()
        if (
            _VOTE_RE.match(content)
            or _MTG_NUM_RE.match(content)
            or _WMPP_RE.match(content)
            or _ISO_STAFF_RE.match(content)
            or _ORG_PRESENTER_RE.match(content)
            or _INIT_CODE_RE.match(content)
        ):
            return True
    return False


def _extract_item_metadata(desc_cell: str, time_cell: str | None = None) -> dict:
    """Extract structured metadata from a description cell.

    Scans all parenthetical groups in the text after the title line, classifies
    each one, then removes them to leave free-text notes.

    Returns a dict (all values may be None / empty list):
        presenter, org, vote_status, meeting_number, wmpp_id,
        initiative_codes, other_tags, notes, time_slot
    """
    lines = desc_cell.strip().splitlines()
    # Collect lines after the title, stopping before any embedded "a) …" sub-item line
    # so the parent item doesn't inherit the sub-item's presenter/metadata.
    rest_lines: list[str] = []
    for line in lines[1:]:
        if _DESC_SUB_LETTER_RE.match(line.strip()):
            break
        rest_lines.append(line)
    rest = "\n".join(rest_lines).strip()

    presenter: str | None = None
    org: str | None = None
    vote_status: str | None = None
    meeting_number: str | None = None
    wmpp_id: str | None = None
    initiative_codes: list[str] = []
    other_tags: list[str] = []
    matched_spans: list[tuple[int, int]] = []

    for m in _PAREN_RE.finditer(rest):
        content = m.group(1).strip()
        matched_spans.append(m.span())

        # Check specific patterns first (most distinctive), presenter last
        # so the broad colon-matcher doesn't swallow WMPP IDs etc.

        # Vote / status tags
        if _VOTE_RE.match(content):
            vote_status = content
            continue

        # Meeting ordinal
        mtg_m = _MTG_NUM_RE.match(content)
        if mtg_m:
            meeting_number = mtg_m.group(1)
            continue

        # WMPP ID
        wmpp_m = _WMPP_RE.match(content)
        if wmpp_m:
            wmpp_id = wmpp_m.group(1)
            continue

        # Initiative code (e.g. CAR-SA, GISWG)
        if _INIT_CODE_RE.match(content):
            initiative_codes.append(content)
            continue

        # Presenter — "(ISO Staff)" or "(OrgName: Names)" for any organisation
        if _ISO_STAFF_RE.match(content):
            org = "ISO-NE" if "ne" in content.lower() else "ISO"
            presenter = "ISO Staff"
            continue

        op_m = _ORG_PRESENTER_RE.match(content)
        if op_m:
            org = op_m.group(1).strip()
            presenter = op_m.group(2).strip()
            continue

        other_tags.append(content)

    # Notes = rest with all matched parenthetical spans blanked out
    notes_chars = list(rest)
    for start, end in matched_spans:
        for i in range(start, end):
            notes_chars[i] = " "
    notes = re.sub(r"\s+", " ", "".join(notes_chars)).strip() or None

    return {
        "presenter":        presenter,
        "org":              org,
        "vote_status":      vote_status,
        "meeting_number":   meeting_number,
        "wmpp_id":          wmpp_id,
        "initiative_codes": initiative_codes,
        "other_tags":       other_tags,
        "notes":            notes,
        "time_slot":        time_cell.strip() if time_cell else None,
    }

# ---------------------------------------------------------------------------
# Table parser
# ---------------------------------------------------------------------------

def _parse_agenda_from_tables(doc: DocxDocument) -> list[dict]:
    """
    Walk all tables in the docx.  For each row whose first cell matches
    _ITEM_ID_RE, emit an agenda item dict:
      {item_id, title, prefix, auto_sub}

    Sub-letter extraction rules:
    1. If the description cell starts with "a)" / "b)" etc., the letter is
       folded into the item_id (e.g. raw_id="4.1", desc="b) Gas Capacity..."
       → item_id="4.1.b").
    2. Duplicate item_ids (no description sub-letter) are auto-assigned .a/.b/…
       in the _assign_auto_sub_prefixes post-processing step.
    """
    seen: set[tuple[str, str]] = set()
    items: list[dict] = []

    for table in doc.tables:
        last_item_id: str | None = None   # tracks last valid item_id within this table
        sub_ctr: dict[str, int] = {}      # counts empty-ID sub-items per parent

        for row in table.rows:
            cells = row.cells
            if not cells:
                continue
            id_cell = cells[0].text.strip()
            desc_cell_raw = cells[1].text.strip() if len(cells) > 1 else ""

            # ── Empty-ID sub-item rows (RC-style v3) ──────────────────────────
            # Some RC agendas place a sub-topic on a separate row with an empty
            # first cell but no "a)" prefix.  Detect them by checking for at least
            # one parenthetical matching a known metadata pattern.
            if not id_cell:
                if last_item_id and _has_metadata_parens(desc_cell_raw):
                    n = sub_ctr.get(last_item_id, 0)
                    sub_ctr[last_item_id] = n + 1
                    sub_letter = chr(ord("a") + n)
                    sub_id = f"{last_item_id}.{sub_letter}"

                    first_line = desc_cell_raw.splitlines()[0].strip()
                    sub_title_raw = re.split(r"\s*\(", first_line)[0].strip().rstrip(".;:*")
                    if not sub_title_raw:
                        continue

                    sub_key = (sub_id, sub_title_raw)
                    if sub_key in seen:
                        continue
                    seen.add(sub_key)

                    # Build synthetic cell: title on line 0, remainder on line 1+
                    parens_part = first_line[len(sub_title_raw):]
                    extra_lines = "\n".join(desc_cell_raw.splitlines()[1:])
                    sub_meta = _extract_item_metadata(
                        f"{sub_title_raw}\n{parens_part}\n{extra_lines}".strip(),
                        None,
                    )
                    items.append({
                        "item_id": sub_id,
                        "title":   sub_title_raw,
                        "prefix":  item_id_to_prefix(sub_id),
                        "auto_sub": False,
                        **sub_meta,
                    })
                continue

            if not _ITEM_ID_RE.match(id_cell):
                continue

            raw_id = id_cell.rstrip("*").strip()
            if raw_id.endswith(".0"):
                raw_id = raw_id[:-2]

            desc_cell = cells[1].text.strip() if len(cells) > 1 else id_cell
            title_line = desc_cell.splitlines()[0].strip().rstrip(".;:*")

            time_cell = cells[2].text.strip() if len(cells) > 2 else None

            sub_m = _DESC_SUB_LETTER_RE.match(title_line)
            if sub_m:
                sub_letter = sub_m.group(1).lower()
                sub_rest   = sub_m.group(2).strip()  # everything after "a) "
                first_line = sub_rest.splitlines()[0].strip()
                # Title = text before the first inline parenthetical on that line
                title_raw  = re.split(r"\s*\(", first_line)[0].strip().rstrip(".;:*")
                title      = title_raw or first_line.rstrip(".;:*")
                effective_id = f"{raw_id}.{sub_letter}"
                # Build synthetic cell: clean title on line 0, inline parens on
                # line 1, any following lines (notes) appended after.
                inline_parens = first_line[len(title_raw):]
                extra_lines   = "\n".join(sub_rest.splitlines()[1:])
                meta_cell     = f"{title}\n{inline_parens}\n{extra_lines}".strip()
            else:
                title        = title_line
                effective_id = raw_id
                meta_cell    = desc_cell  # existing behaviour

            if not title:
                continue
            key = (effective_id, title)
            if key in seen:
                continue
            seen.add(key)

            # Drop "Continued" duplicates: same item_id, title ends with ", Continued"
            base_title = re.sub(r",?\s*continued\.?$", "", title, flags=re.IGNORECASE).strip()
            if any(
                it["item_id"] == effective_id and
                it["title"].lower() == base_title.lower()
                for it in items
            ):
                continue

            # Extract metadata from the (possibly synthetic) meta cell
            meta = _extract_item_metadata(meta_cell, time_cell)

            items.append({
                "item_id": effective_id,
                "title": base_title if base_title != title else title,
                "prefix": item_id_to_prefix(effective_id),
                "auto_sub": False,
                **meta,
            })

            # Track for empty-ID sub-item detection (only update on non-sub-letter rows
            # so that "4.1.a", "4.1.b" don't reset the parent to "4.1.a")
            if not sub_m:
                last_item_id = effective_id

            # Detect embedded sub-items: lines inside this cell that start with
            # "a) Title (Presenter) ..." — these appear in RC-style agendas where
            # sub-items are body text rather than separate table rows.
            # Only emit them when the title_line itself is NOT already a sub-item
            # (avoid double-counting when description starts with "a) …").
            if not sub_m:
                for body_line in desc_cell.splitlines()[1:]:
                    body_stripped = body_line.strip()
                    emb_m = _DESC_SUB_LETTER_RE.match(body_stripped)
                    if not emb_m:
                        continue
                    sub_letter = emb_m.group(1).lower()
                    sub_rest   = emb_m.group(2).strip()
                    # Title = text before first '('
                    sub_title_raw = re.split(r"\s*\(", sub_rest)[0].strip().rstrip(".;:*")
                    if not sub_title_raw:
                        continue
                    sub_id = f"{effective_id}.{sub_letter}"
                    sub_key = (sub_id, sub_title_raw)
                    if sub_key in seen:
                        continue
                    seen.add(sub_key)
                    # Build a synthetic cell: title on line 0, parens remainder on line 1
                    parens_part = sub_rest[len(sub_title_raw):]
                    sub_meta = _extract_item_metadata(
                        f"{sub_title_raw}\n{parens_part}", time_cell
                    )
                    items.append({
                        "item_id": sub_id,
                        "title":   sub_title_raw,
                        "prefix":  item_id_to_prefix(sub_id),
                        "auto_sub": False,
                        **sub_meta,
                    })

    return items


def _assign_auto_sub_prefixes(items: list[dict]) -> list[dict]:
    """
    When multiple agenda rows share the same item_id (no sub-letter in the
    agenda table), auto-assign sequential letters a, b, c… so each gets a
    unique item_id and prefix.

    The first occurrence keeps the bare id (e.g. "7.1") with auto_sub=False
    only if there is a file with that bare prefix (handled at display time).
    Actually we assign .a to the first occurrence as well so every item has
    a unique key.  The bare prefix ("a07.1") is treated as an alias for .a
    in map_docs_to_agenda_items.
    """
    from collections import Counter
    counts = Counter(i["item_id"] for i in items)
    seq: dict[str, int] = {}
    result = []
    for item in items:
        iid = item["item_id"]
        if counts[iid] > 1:
            n = seq.get(iid, 0)
            seq[iid] = n + 1
            letter = chr(ord("a") + n)
            new_id = f"{iid}.{letter}"
            base_pfx = item.get("prefix") or ""
            new_pfx = f"{base_pfx}.{letter}" if base_pfx else None
            result.append({**item, "item_id": new_id, "prefix": new_pfx, "auto_sub": True})
        else:
            result.append(item)
    return result


def _drop_continued(items: list[dict]) -> list[dict]:
    """
    Remove agenda rows whose title ends with ', Continued' (case-insensitive)
    when a row with the same item_id and the base title already exists in the list.
    These are lunch-break continuations with no additional files.
    """
    _CONT_RE = re.compile(r",?\s*continued\.?$", re.IGNORECASE)
    base_keys: set[tuple[str, str]] = set()
    for item in items:
        if not _CONT_RE.search(item["title"]):
            base_keys.add((item["item_id"], _CONT_RE.sub("", item["title"]).strip().lower()))

    result = []
    for item in items:
        if _CONT_RE.search(item["title"]):
            base = _CONT_RE.sub("", item["title"]).strip().lower()
            if (item["item_id"], base) in base_keys:
                continue  # skip the Continued duplicate
        result.append(item)
    return result


# ---------------------------------------------------------------------------
# Public: parse a docx file → list of agenda items
# ---------------------------------------------------------------------------

def parse_agenda_from_docx(path_or_bytes) -> list[dict]:
    """
    Parse an agenda .docx and return a list of agenda item dicts, each:
      {item_id, title, prefix, auto_sub}

    `path_or_bytes` may be a file path (str/Path) or raw bytes.
    """
    if isinstance(path_or_bytes, (str, Path)):
        doc = DocxDocument(str(path_or_bytes))
    else:
        import io
        doc = DocxDocument(io.BytesIO(path_or_bytes))

    items = _parse_agenda_from_tables(doc)
    items = _drop_continued(items)
    items = _assign_auto_sub_prefixes(items)
    return items


# ---------------------------------------------------------------------------
# Public: map filenames → agenda items
# ---------------------------------------------------------------------------

def _prefix_matches(filename_lower: str, prefix: str) -> bool:
    """True if filename starts with prefix followed by a separator or end."""
    if not filename_lower.startswith(prefix):
        return False
    rest = filename_lower[len(prefix):]
    return not rest or rest[0] in ("_", ".", "-")


def map_docs_to_agenda_items(
    doc_rows: list[dict],
    agenda_items: list[dict],
) -> dict[str, list]:
    """
    Assign each document row to the best-matching agenda item bucket.

    Bucket key = item prefix (e.g. "a07.1.b").
    Special rule: for auto-sub .a items, the bare base prefix (e.g. "a07.1")
    is treated as an alias so files named a07.1_foo.pdf map to the .a bucket.

    Returns {prefix: [doc_row, ...], "other": [...]} — "other" holds unmatched.
    """
    # Build match entries: (match_prefix, bucket_key)
    match_entries: list[tuple[str, str]] = []
    buckets: dict[str, list] = {"other": []}

    for item in agenda_items:
        p = item.get("prefix")
        if not p:
            continue
        buckets[p] = []
        match_entries.append((p, p))
        # For auto-sub .a items, also match the bare base prefix
        if item.get("auto_sub") and p.endswith(".a"):
            base = p[:-2]  # "a07.1.a" → "a07.1"
            # Only add if no item already claims the bare prefix as its own
            if not any(i.get("prefix") == base for i in agenda_items):
                match_entries.append((base, p))

    for row in doc_rows:
        fn = row.get("filename", "").lower()
        best_bucket = "other"
        best_len = 0
        for match_pfx, bucket_key in match_entries:
            if _prefix_matches(fn, match_pfx) and len(match_pfx) > best_len:
                best_bucket = bucket_key
                best_len = len(match_pfx)
        buckets[best_bucket].append(row)

    return {k: v for k, v in buckets.items() if v}
